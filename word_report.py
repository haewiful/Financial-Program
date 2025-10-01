import openpyxl
from docx import Document
import os
import datetime

class WordReportGenerator:
    """Reads data from an Excel file and generates a Word report (.docx)."""

    def generate_report(self, excel_file_path):
        """Reads data from the Excel path and saves the Word report."""
        
        if not os.path.exists(excel_file_path):
            raise FileNotFoundError(f"File not found: {excel_file_path}")
        print(f"file found: {excel_file_path}")

        # The output file will be in the same directory as the input file
        base_dir = os.path.dirname(excel_file_path)
        base_name = os.path.basename(excel_file_path)
        doc_name = os.path.join(base_dir, base_name.replace('.xlsx', '_Report.docx'))

        print("before workebook")

        workbook = openpyxl.load_workbook(excel_file_path)
        print("after workbook")
        sheet = workbook.active
        print("active")
        
        headers = [cell.value for cell in sheet[1]]
        data_rows = list(sheet.iter_rows(min_row=2, values_only=True)) 

        if not data_rows:
            raise ValueError("The Excel file is empty or only contains headers.")

        # Create the Word Document
        document = Document()
        print("Document()")
        document.add_heading(f'Report Generated from: {base_name}', 0)
        print("add_heading")
        document.add_paragraph(f"Report Date: {datetime.datetime.now().strftime('%Y-%m-%d')}")
        print("add_paragraph")
        document.add_heading('Data Summary Table', level=1)

        print("document")

        # Create a Table in Word
        table = document.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'

        print("add table")

        # Set the table headers
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = str(header)
        
        print("header")
            
        # Populate the table with Excel data
        for row_data in data_rows:
            row_cells = table.add_row().cells
            for i, cell_value in enumerate(row_data):
                row_cells[i].text = str(cell_value)
        
        print("data writing")
        
        # Save the Word File
        document.save(doc_name)
        print("save")
        # Return the save path for the GUI to display
        return doc_name